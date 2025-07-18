from docx import Document

# Função para criar o Word
def cria_arquivo_word(mensagens, nome_arquivo='mensagens.docx'):
    doc = Document()
    
    doc.add_heading('Relatório de Mensagens', 0)

    for idx, msg in enumerate(mensagens, 1):
        doc.add_heading(f'Mensagem {idx}', level=1)
        doc.add_paragraph(msg)

    doc.save(nome_arquivo)
    print(f'Arquivo "{nome_arquivo}" criado com sucesso!')

# Exemplo de mensagens capturadas
mensagens = [
    "E aí, beleza?",
    "Tô te mandando as mensagens.",
    "Depois gera o Word!",
    "Fechou!"
]

cria_arquivo_word(mensagens)